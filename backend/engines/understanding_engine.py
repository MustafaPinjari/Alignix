"""
Document Understanding Engine — deep structural analysis with confidence scoring.
Every detected element carries a confidence score and correction risk level.
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import Counter
import re

ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    None: "left",
}

# Confidence thresholds
HIGH_CONF   = 0.85
MEDIUM_CONF = 0.60
LOW_CONF    = 0.40


class DocumentUnderstandingEngine:

    def understand(self, path: str) -> dict:
        doc = Document(path)
        raw_paras = list(doc.paragraphs)
        fonts      = self._collect_fonts(raw_paras)
        dom_font   = Counter(fonts).most_common(1)[0][0] if fonts else None
        body_sizes = self._body_sizes(raw_paras)
        dom_size   = Counter(body_sizes).most_common(1)[0][0] if body_sizes else 12.0

        elements = [
            self._analyze_element(i, p, raw_paras, dom_font, dom_size)
            for i, p in enumerate(raw_paras)
        ]
        tables = self._analyze_tables(doc)

        structure_conf  = self._structure_confidence(elements)
        layout_conf     = self._layout_confidence(elements, tables)
        correction_risk = self._correction_risk(elements, tables)

        clarifications = self._build_clarifications(elements)
        issues         = self._build_issues(elements, tables, dom_font, dom_size)

        return {
            "elements": elements,
            "tables": tables,
            "dominant_font": dom_font,
            "dominant_size": dom_size,
            "confidence": {
                "structure": round(structure_conf, 2),
                "layout":    round(layout_conf, 2),
                "correction_risk": round(correction_risk, 2),
            },
            "clarifications": clarifications,
            "issues": issues,
            "stats": self._stats(elements, tables),
        }

    # ── Element Analysis ──────────────────────────────────────────────────────

    def _analyze_element(self, idx, para, all_paras, dom_font, dom_size) -> dict:
        text       = para.text.strip()
        style      = para.style.name if para.style else "Normal"
        style_low  = style.lower()
        runs       = para.runs
        font_sizes = [r.font.size.pt for r in runs if r.font.size]
        font_names = [r.font.name   for r in runs if r.font.name]
        bold_flags = [r.bold        for r in runs if r.bold is not None]
        size       = max(font_sizes) if font_sizes else None
        bold       = any(bold_flags)
        font_name  = font_names[0] if font_names else None

        role, role_conf = self._infer_role(idx, text, style_low, size, bold, dom_size, all_paras)
        issues          = self._element_issues(text, style_low, size, bold, font_name, dom_font, dom_size, role)
        risk            = self._element_risk(role, role_conf, issues)

        return {
            "index":       idx,
            "text":        text[:150],
            "style":       style,
            "role":        role,
            "confidence":  round(role_conf, 2),
            "risk":        risk,
            "font_size":   round(size, 1) if size else None,
            "font_name":   font_name,
            "bold":        bold,
            "alignment":   ALIGN_MAP.get(para.alignment, "left"),
            "line_spacing": self._line_spacing(para),
            "space_before": para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 0,
            "space_after":  para.paragraph_format.space_after.pt  if para.paragraph_format.space_after  else 0,
            "issues":      issues,
            "empty":       not bool(text),
        }

    def _infer_role(self, idx, text, style_low, size, bold, dom_size, all_paras):
        # Style-based — highest confidence
        if "heading 1" in style_low: return "heading1", 0.97
        if "heading 2" in style_low: return "heading2", 0.97
        if "heading 3" in style_low: return "heading3", 0.95
        if "heading"   in style_low: return "heading",  0.90
        if "title"     in style_low: return "title",    0.95
        if "caption"   in style_low: return "caption",  0.92
        if "list"      in style_low or "bullet" in style_low: return "list_item", 0.88

        if not text:
            return "empty", 1.0

        # Pattern-based
        if re.match(r"^(abstract|introduction|conclusion|references|bibliography)$", text, re.I):
            return "chapter", 0.88
        if re.match(r"^appendix\s+[a-z]", text, re.I):
            return "appendix", 0.85
        if re.match(r"^\[\d+\]", text):
            return "reference", 0.82
        if re.match(r"^(table|figure|fig\.|tab\.)\s*\d+", text, re.I):
            return "caption", 0.80
        if re.match(r"^\d+\.\d+\.\d+\s+\S", text):
            return "heading3", 0.78
        if re.match(r"^\d+\.\d+\s+\S", text):
            return "heading2", 0.75
        if re.match(r"^\d+\.\s+[A-Z]", text):
            return "heading1", 0.72

        # Size + bold heuristics
        if size and dom_size:
            if bold and size >= dom_size + 6: return "heading1", 0.70
            if bold and size >= dom_size + 3: return "heading2", 0.65
            if bold and size >= dom_size + 1: return "heading3", 0.58
            if bold and len(text) < 80:       return "heading3", 0.50

        # Short all-caps line
        if text.isupper() and len(text) < 60 and len(text) > 3:
            return "heading1", 0.55

        return "body", 0.90

    def _element_issues(self, text, style_low, size, bold, font_name, dom_font, dom_size, role) -> list:
        issues = []
        if font_name and dom_font and font_name != dom_font and role == "body":
            issues.append({"type": "font_inconsistency", "detail": f"Expected {dom_font}, found {font_name}"})
        if size and dom_size and role == "body" and abs(size - dom_size) > 2:
            issues.append({"type": "size_inconsistency", "detail": f"Expected {dom_size}pt, found {size}pt"})
        return issues

    def _element_risk(self, role, confidence, issues) -> str:
        if confidence < LOW_CONF:
            return "high"
        if confidence < MEDIUM_CONF or issues:
            return "medium"
        return "low"

    # ── Tables ────────────────────────────────────────────────────────────────

    def _analyze_tables(self, doc) -> list:
        results = []
        for i, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns)
            # Estimate overflow risk by column count
            overflow_risk = "high" if cols > 6 else "medium" if cols > 4 else "low"
            results.append({
                "index": i, "rows": rows, "cols": cols,
                "overflow_risk": overflow_risk,
                "confidence": 0.95,
            })
        return results

    # ── Confidence Scores ─────────────────────────────────────────────────────

    def _structure_confidence(self, elements) -> float:
        scored = [e for e in elements if not e["empty"]]
        if not scored:
            return 0.5
        return sum(e["confidence"] for e in scored) / len(scored)

    def _layout_confidence(self, elements, tables) -> float:
        risks = [e for e in elements if e["risk"] == "high"]
        high_table_risks = [t for t in tables if t["overflow_risk"] == "high"]
        total = len(elements) + len(tables) + 1
        risk_count = len(risks) + len(high_table_risks)
        return max(0.1, 1.0 - (risk_count / total))

    def _correction_risk(self, elements, tables) -> float:
        low_conf = [e for e in elements if e["confidence"] < MEDIUM_CONF and not e["empty"]]
        total = len([e for e in elements if not e["empty"]]) + 1
        return min(1.0, len(low_conf) / total)

    # ── Clarifications ────────────────────────────────────────────────────────

    def _build_clarifications(self, elements) -> list:
        """Generate clarification questions only for low-confidence elements."""
        questions = []
        seen_types = set()

        for el in elements:
            if el["empty"] or el["confidence"] >= MEDIUM_CONF:
                continue
            key = (el["role"], round(el["confidence"], 1))
            if key in seen_types:
                continue
            seen_types.add(key)

            q = self._make_question(el)
            if q:
                questions.append(q)
            if len(questions) >= 5:  # cap at 5 questions max
                break

        return questions

    def _make_question(self, el) -> dict | None:
        role = el["role"]
        text = el["text"][:80]
        conf_pct = round(el["confidence"] * 100)

        if role in ("heading1", "heading2", "heading3", "heading"):
            return {
                "id": f"q_role_{el['index']}",
                "element_index": el["index"],
                "type": "role_clarification",
                "confidence": el["confidence"],
                "question": f'This text appears to be a heading ({conf_pct}% confidence):\n"{text}"',
                "options": [
                    {"value": "heading1", "label": "Heading 1"},
                    {"value": "heading2", "label": "Heading 2"},
                    {"value": "heading3", "label": "Heading 3"},
                    {"value": "body",     "label": "Normal Body Text"},
                ],
                "current": role,
            }
        if role == "body" and el.get("issues"):
            issue = el["issues"][0]
            if issue["type"] == "font_inconsistency":
                return {
                    "id": f"q_font_{el['index']}",
                    "element_index": el["index"],
                    "type": "font_clarification",
                    "confidence": el["confidence"],
                    "question": f'Multiple font styles detected near:\n"{text}"',
                    "options": [
                        {"value": "normalize", "label": "Normalize to dominant font"},
                        {"value": "keep",      "label": "Keep original"},
                    ],
                    "current": "keep",
                }
        return None

    # ── Issues ────────────────────────────────────────────────────────────────

    def _build_issues(self, elements, tables, dom_font, dom_size) -> list:
        issues = []
        for el in elements:
            for iss in el.get("issues", []):
                issues.append({
                    **iss,
                    "element_index": el["index"],
                    "element_role":  el["role"],
                    "text":          el["text"][:80],
                    "severity":      "high" if el["risk"] == "high" else "medium",
                    "confidence":    el["confidence"],
                })
        # Orphan heading detection
        roles = [e["role"] for e in elements if not e["empty"]]
        for i, el in enumerate(elements):
            if "heading" in el["role"] and not el["empty"]:
                # heading at end of document
                remaining = [e for e in elements[i+1:] if not e["empty"]]
                if not remaining:
                    issues.append({
                        "type": "orphan_heading",
                        "element_index": el["index"],
                        "element_role": el["role"],
                        "text": el["text"][:80],
                        "severity": "medium",
                        "detail": "Heading at end of document with no following content",
                        "confidence": 0.9,
                    })
        return issues

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _collect_fonts(self, paras) -> list:
        return [r.font.name for p in paras for r in p.runs if r.font.name]

    def _body_sizes(self, paras) -> list:
        sizes = []
        for p in paras:
            style = p.style.name.lower() if p.style else ""
            if "heading" in style or "title" in style:
                continue
            for r in p.runs:
                if r.font.size:
                    sizes.append(round(r.font.size.pt, 1))
        return sizes

    def _line_spacing(self, para) -> float:
        fmt = para.paragraph_format
        if fmt.line_spacing is None:
            return 1.0
        if isinstance(fmt.line_spacing, int):
            return round(fmt.line_spacing / 240, 2)
        try:
            return round(fmt.line_spacing.pt / 12, 2)
        except Exception:
            return 1.0

    def _stats(self, elements, tables) -> dict:
        non_empty = [e for e in elements if not e["empty"]]
        return {
            "total_paragraphs": len(non_empty),
            "headings":  len([e for e in non_empty if "heading" in e["role"]]),
            "body":      len([e for e in non_empty if e["role"] == "body"]),
            "tables":    len(tables),
            "lists":     len([e for e in non_empty if e["role"] == "list_item"]),
            "captions":  len([e for e in non_empty if e["role"] == "caption"]),
            "references":len([e for e in non_empty if e["role"] == "reference"]),
            "low_confidence": len([e for e in non_empty if e["confidence"] < MEDIUM_CONF]),
        }
