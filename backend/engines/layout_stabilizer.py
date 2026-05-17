"""
Layout Stabilization Engine — prevents broken pagination, orphan headings,
table overflow, and spacing collapse after formatting corrections.
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree


class LayoutStabilizer:

    def stabilize(self, doc: Document) -> list[str]:
        """Run all stabilization passes. Returns list of actions taken."""
        actions = []
        actions += self._fix_orphan_headings(doc)
        actions += self._fix_table_overflow(doc)
        actions += self._fix_spacing_collapse(doc)
        actions += self._fix_widow_control(doc)
        return actions

    # ── Orphan Heading Prevention ─────────────────────────────────────────────

    def _fix_orphan_headings(self, doc: Document) -> list[str]:
        """Ensure headings have keep-with-next set so they don't orphan."""
        actions = []
        for para in doc.paragraphs:
            style = para.style.name.lower() if para.style else ""
            if "heading" in style:
                pPr = para._p.get_or_add_pPr()
                kwn = pPr.find(qn("w:keepNext"))
                if kwn is None:
                    kwn = etree.SubElement(pPr, qn("w:keepNext"))
                    actions.append(f"keep-with-next: {para.text[:40]}")
        return actions

    # ── Table Overflow Prevention ─────────────────────────────────────────────

    def _fix_table_overflow(self, doc: Document) -> list[str]:
        """Set tables to autofit to prevent overflow beyond page margins."""
        actions = []
        for i, table in enumerate(doc.tables):
            tbl_pr = table._tbl.find(qn("w:tblPr"))
            if tbl_pr is None:
                tbl_pr = etree.SubElement(table._tbl, qn("w:tblPr"))
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = etree.SubElement(tbl_pr, qn("w:tblLayout"))
            if layout.get(qn("w:type")) != "autofit":
                layout.set(qn("w:type"), "autofit")
                actions.append(f"table[{i}] autofit")
        return actions

    # ── Spacing Collapse Prevention ───────────────────────────────────────────

    def _fix_spacing_collapse(self, doc: Document) -> list[str]:
        """Ensure body paragraphs have minimum space_after to prevent collapse."""
        actions = []
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            style = para.style.name.lower() if para.style else ""
            if "heading" in style or "title" in style:
                continue
            fmt = para.paragraph_format
            after = fmt.space_after
            if after is not None and after.pt == 0:
                fmt.space_after = Pt(4)
                actions.append(f"spacing-collapse fix: {para.text[:30]}")
        return actions

    # ── Widow/Orphan Control ──────────────────────────────────────────────────

    def _fix_widow_control(self, doc: Document) -> list[str]:
        """Enable widow/orphan control on all body paragraphs."""
        actions = []
        for para in doc.paragraphs:
            style = para.style.name.lower() if para.style else ""
            if "heading" in style:
                continue
            pPr = para._p.get_or_add_pPr()
            wc = pPr.find(qn("w:widowControl"))
            if wc is None:
                wc = etree.SubElement(pPr, qn("w:widowControl"))
                wc.set(qn("w:val"), "1")
                actions.append(f"widow-control: {para.text[:30]}")
        return actions
