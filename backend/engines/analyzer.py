"""
Document Analyzer — heuristic structural analysis of DOCX files.
Detects headings, body, tables, lists, spacing, font inconsistencies.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import Counter
import re


ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    None: "left",
}


class DocumentAnalyzer:

    def analyze(self, path: str) -> dict:
        doc = Document(path)
        paragraphs = self._analyze_paragraphs(doc)
        tables = self._analyze_tables(doc)
        fonts = self._collect_fonts(doc)
        issues = self._detect_issues(paragraphs, fonts)
        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "fonts": fonts,
            "issues": issues,
            "stats": self._stats(paragraphs, tables),
        }

    def health_score(self, path: str) -> dict:
        data = self.analyze(path)
        issues = data["issues"]
        total = len(data["paragraphs"]) + len(data["tables"]) + 1
        issue_count = len(issues)
        integrity = max(0, round(100 - (issue_count / total) * 100, 1))
        font_variety = len(set(data["fonts"]))
        professionalism = max(0, round(100 - font_variety * 5, 1))
        readability = self._readability_score(data["paragraphs"])
        structural = self._structural_score(data["paragraphs"])
        return {
            "integrity": integrity,
            "professionalism": professionalism,
            "readability": readability,
            "structural": structural,
            "overall": round((integrity + professionalism + readability + structural) / 4, 1),
        }

    # ── Internal ──────────────────────────────────────────────────────────────

    def _analyze_paragraphs(self, doc) -> list:
        results = []
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
            style_name = para.style.name if para.style else "Normal"
            font_sizes = [r.font.size for r in para.runs if r.font.size]
            font_names = [r.font.name for r in para.runs if r.font.name]
            bold_flags = [r.bold for r in para.runs if r.bold is not None]
            results.append({
                "index": i,
                "text": para.text[:120],
                "style": style_name,
                "element_type": self._classify_element(style_name, para.text),
                "font_size": round(font_sizes[0].pt, 1) if font_sizes else None,
                "font_name": font_names[0] if font_names else None,
                "bold": any(bold_flags),
                "alignment": ALIGN_MAP.get(para.alignment, "left"),
                "line_spacing": self._get_line_spacing(para),
                "space_before": para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 0,
                "space_after": para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 0,
            })
        return results

    def _analyze_tables(self, doc) -> list:
        results = []
        for i, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns)
            results.append({
                "index": i,
                "rows": rows,
                "cols": cols,
                "alignment": self._get_table_alignment(table),
            })
        return results

    def _collect_fonts(self, doc) -> list:
        fonts = []
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    fonts.append(run.font.name)
        return fonts

    def _detect_issues(self, paragraphs: list, fonts: list) -> list:
        issues = []
        font_counter = Counter(fonts)
        dominant_font = font_counter.most_common(1)[0][0] if font_counter else None

        sizes = [p["font_size"] for p in paragraphs if p["font_size"] and p["element_type"] == "body"]
        dominant_size = Counter(sizes).most_common(1)[0][0] if sizes else None

        for p in paragraphs:
            if dominant_font and p["font_name"] and p["font_name"] != dominant_font:
                issues.append({"index": p["index"], "type": "font_inconsistency",
                                "detail": f"Expected {dominant_font}, found {p['font_name']}"})
            if p["element_type"] == "body" and dominant_size and p["font_size"] and abs(p["font_size"] - dominant_size) > 2:
                issues.append({"index": p["index"], "type": "size_inconsistency",
                                "detail": f"Expected {dominant_size}pt, found {p['font_size']}pt"})
            if p["element_type"] == "body" and p["alignment"] not in ("justify", "left"):
                issues.append({"index": p["index"], "type": "alignment_issue",
                                "detail": f"Body paragraph aligned {p['alignment']}"})
        return issues

    def _classify_element(self, style_name: str, text: str) -> str:
        s = style_name.lower()
        if "heading 1" in s: return "heading1"
        if "heading 2" in s: return "heading2"
        if "heading 3" in s: return "heading3"
        if "heading" in s: return "heading"
        if "caption" in s: return "caption"
        if "list" in s: return "list"
        if "title" in s: return "title"
        return "body"

    def _get_line_spacing(self, para) -> float:
        fmt = para.paragraph_format
        if fmt.line_spacing is None:
            return 1.0
        if isinstance(fmt.line_spacing, int):
            return round(fmt.line_spacing / 240, 2)
        try:
            return round(fmt.line_spacing.pt / 12, 2)
        except Exception:
            return 1.0

    def _get_table_alignment(self, table) -> str:
        try:
            from docx.oxml.ns import qn
            tbl_pr = table._tbl.find(qn("w:tblPr"))
            if tbl_pr is not None:
                jc = tbl_pr.find(qn("w:jc"))
                if jc is not None:
                    return jc.get(qn("w:val"), "left")
        except Exception:
            pass
        return "left"

    def _readability_score(self, paragraphs: list) -> float:
        body = [p for p in paragraphs if p["element_type"] == "body"]
        if not body:
            return 100.0
        avg_len = sum(len(p["text"]) for p in body) / len(body)
        score = max(0, 100 - max(0, avg_len - 200) * 0.2)
        return round(score, 1)

    def _structural_score(self, paragraphs: list) -> float:
        headings = [p for p in paragraphs if "heading" in p["element_type"]]
        body = [p for p in paragraphs if p["element_type"] == "body"]
        if not headings:
            return 60.0
        ratio = len(body) / max(len(headings), 1)
        score = min(100, 60 + ratio * 5)
        return round(score, 1)

    def _stats(self, paragraphs: list, tables: list) -> dict:
        return {
            "total_paragraphs": len(paragraphs),
            "headings": len([p for p in paragraphs if "heading" in p["element_type"]]),
            "body_paragraphs": len([p for p in paragraphs if p["element_type"] == "body"]),
            "tables": len(tables),
            "lists": len([p for p in paragraphs if p["element_type"] == "list"]),
        }
