"""
Formatting Rule Engine — applies profile rules to document elements.
Supports heading1-3, body, caption, list, table element types.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import copy

ALIGN_ENUM = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class RuleEngine:

    def apply_rules(self, doc: Document, rules: list) -> tuple[Document, list]:
        """Apply all rules to document. Returns (modified_doc, change_log)."""
        rule_map = {r["element"]: r for r in rules}
        log = []

        for para in doc.paragraphs:
            element_type = self._classify(para)
            rule = rule_map.get(element_type)
            if not rule:
                continue
            changes = self._apply_paragraph_rule(para, rule)
            if changes:
                log.append({"element": element_type, "text": para.text[:60], "changes": changes})

        if "table" in rule_map:
            for table in doc.tables:
                self._apply_table_rule(table, rule_map["table"])

        return doc, log

    def check_violations(self, doc: Document, rules: list) -> list:
        """Return list of violations without modifying document."""
        rule_map = {r["element"]: r for r in rules}
        violations = []

        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
            element_type = self._classify(para)
            rule = rule_map.get(element_type)
            if not rule:
                continue
            v = self._check_paragraph_violations(i, para, rule)
            violations.extend(v)

        return violations

    # ── Internal ──────────────────────────────────────────────────────────────

    def _classify(self, para) -> str:
        s = para.style.name.lower() if para.style else "normal"
        if "heading 1" in s: return "heading1"
        if "heading 2" in s: return "heading2"
        if "heading 3" in s: return "heading3"
        if "heading" in s: return "heading1"
        if "caption" in s: return "caption"
        if "list" in s: return "list"
        return "body"

    def _apply_paragraph_rule(self, para, rule: dict) -> list:
        changes = []

        if rule.get("alignment"):
            target = ALIGN_ENUM.get(rule["alignment"])
            if target is not None and para.alignment != target:
                para.alignment = target
                changes.append(f"alignment→{rule['alignment']}")

        fmt = para.paragraph_format
        if rule.get("line_spacing"):
            from docx.shared import Pt
            target_spacing = rule["line_spacing"]
            fmt.line_spacing = Pt(target_spacing * 12)
            changes.append(f"line_spacing→{target_spacing}")

        if rule.get("space_before") is not None:
            fmt.space_before = Pt(rule["space_before"])
        if rule.get("space_after") is not None:
            fmt.space_after = Pt(rule["space_after"])

        for run in para.runs:
            run_changes = self._apply_run_rule(run, rule)
            changes.extend(run_changes)

        return list(set(changes))

    def _apply_run_rule(self, run, rule: dict) -> list:
        changes = []
        if rule.get("font_name") and run.font.name != rule["font_name"]:
            run.font.name = rule["font_name"]
            changes.append(f"font→{rule['font_name']}")
        if rule.get("font_size"):
            target = Pt(rule["font_size"])
            if run.font.size != target:
                run.font.size = target
                changes.append(f"size→{rule['font_size']}pt")
        if rule.get("bold") is not None:
            run.bold = bool(rule["bold"])
        if rule.get("italic") is not None:
            run.italic = bool(rule["italic"])
        if rule.get("color"):
            run.font.color.rgb = self._parse_color(rule["color"])
            changes.append(f"color→{rule['color']}")
        return changes

    def _apply_table_rule(self, table, rule: dict):
        from docx.oxml.ns import qn
        tbl_pr = table._tbl.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = etree.SubElement(table._tbl, qn("w:tblPr"))
        jc = tbl_pr.find(qn("w:jc"))
        if jc is None:
            jc = etree.SubElement(tbl_pr, qn("w:jc"))
        align = rule.get("alignment", "center")
        jc.set(qn("w:val"), align)

        layout = tbl_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = etree.SubElement(tbl_pr, qn("w:tblLayout"))
        layout.set(qn("w:type"), "autofit")

    def _check_paragraph_violations(self, idx: int, para, rule: dict) -> list:
        violations = []
        for run in para.runs:
            if rule.get("font_size") and run.font.size:
                actual = round(run.font.size.pt, 1)
                if abs(actual - rule["font_size"]) > 0.5:
                    violations.append({
                        "index": idx, "element": self._classify(para),
                        "type": "font_size", "expected": rule["font_size"], "actual": actual,
                        "text": para.text[:60],
                    })
            if rule.get("font_name") and run.font.name and run.font.name != rule["font_name"]:
                violations.append({
                    "index": idx, "element": self._classify(para),
                    "type": "font_name", "expected": rule["font_name"], "actual": run.font.name,
                    "text": para.text[:60],
                })
        return violations

    def _parse_color(self, color_str: str) -> RGBColor:
        color_str = color_str.lstrip("#")
        r, g, b = int(color_str[0:2], 16), int(color_str[2:4], 16), int(color_str[4:6], 16)
        return RGBColor(r, g, b)
